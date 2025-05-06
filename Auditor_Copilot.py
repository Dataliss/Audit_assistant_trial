import os
import re
import io
import tempfile
import pandas as pd
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import gradio as gr
from groq import Groq  # Make sure to install groq package (pip install groq)
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# API Configuration - use environment variables for security
API_KEY = os.environ.get('GROQ_API_KEY')
if not API_KEY:
    raise EnvironmentError("GROQ_API_KEY environment variable not set")

# Initialize API client
groq_client = Groq(api_key=API_KEY)

# Initialize response cache for performance
response_cache = {}

def call_groq_api(prompt, model="llama3-70b-8192", system_message="You are an expert auditor. Provide professional, accurate insights, With citations and references."):
    """Call the Groq API with caching for performance."""
    # Create a unique key for caching based on parameters
    cache_key = f"{prompt[:100]}_{model}_{system_message}"
    
    if cache_key in response_cache:
        return response_cache[cache_key]
    
    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            model=model,
            temperature=0.1,
            max_tokens=4096,
            top_p=1,
        )
        response = chat_completion.choices[0].message.content
        response_cache[cache_key] = response  # Store the response in cache
        return response
    except Exception as e:
        error_msg = f"Error calling Groq API: {str(e)}"
        print(error_msg)
        return error_msg

# Document parsing functions
def parse_pdf(file_obj):
    """Extract text from PDF files."""
    try:
        pdf_reader = PyPDF2.PdfReader(file_obj)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error parsing PDF: {str(e)}")

def parse_docx(file_obj):
    """Extract text from DOCX files."""
    try:
        doc = Document(file_obj)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")

def parse_excel(file_obj):
    """Extract text from Excel files."""
    try:
        df = pd.read_excel(file_obj)
        return df.to_string(index=False)
    except Exception as e:
        raise ValueError(f"Error processing Excel file: {str(e)}")

def parse_csv(file_obj):
    """Extract text from CSV files."""
    try:
        df = pd.read_csv(file_obj)
        return df.to_string(index=False)
    except Exception as e:
        raise ValueError(f"Error processing CSV file: {str(e)}")

def parse_text(file_obj):
    """Extract text from text files."""
    try:
        content = file_obj.read()
        # Handle binary vs string content
        if isinstance(content, bytes):
            return content.decode('utf-8', errors='replace')
        return content
    except Exception as e:
        raise ValueError(f"Error processing text file: {str(e)}")

def process_uploaded_file(file):
    """Process an uploaded file and extract its text."""
    try:
        # Get file extension
        filename = file.name
        extension = os.path.splitext(filename)[1].lower()
                
        # Process file based on its extension
        if extension == '.pdf':
            return parse_pdf(file), filename
        elif extension in ['.xlsx', '.xls']:
            return parse_excel(file), filename
        elif extension == '.csv':
            return parse_csv(file), filename
        elif extension in ['.docx', '.doc']:
            return parse_docx(file), filename
        elif extension == '.txt':
            return parse_text(file), filename
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    except Exception as e:
        raise ValueError(f"Error processing {file.name}: {str(e)}")

def chunk_text(text, max_chunk_size=4000):
    """Split text into manageable chunks for API processing."""
    chunks = []
    current_chunk = ""
    
    for paragraph in text.split("\n"):
        if len(current_chunk) + len(paragraph) < max_chunk_size:
            current_chunk += paragraph + "\n"
        else:
            chunks.append(current_chunk)
            current_chunk = paragraph + "\n"
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks

def analyze_document(filename, text, audit_type=None):
    """Analyze document content for audit insights."""
    chunks = chunk_text(text, max_chunk_size=3500)
    doc_analysis_results = []
    
    for i, chunk in enumerate(chunks):
        prompt = (
            f"Analyze this section {i+1}/{len(chunks)} of a "
            f"{audit_type or 'Financial Statement Audit'} document:\n\n{chunk}\n\n"
            "Provide key observations, potential risks, and compliance issues."
        )
        
        result = call_groq_api(prompt)
        doc_analysis_results.append(result)
    
    # If multiple chunks, create a summary
    if len(doc_analysis_results) > 1:
        summary_prompt = (
            "Synthesize the following section analyses into a cohesive document analysis:\n\n" + 
            "\n\n".join([f"Section {i+1}:\n{result}" for i, result in enumerate(doc_analysis_results)])
        )
        return call_groq_api(summary_prompt)
    elif doc_analysis_results:
        return doc_analysis_results[0]
    else:
        return "No content to analyze."

def determine_audit_framework(document_texts, audit_type=None):
    """Identify the most appropriate audit framework based on document content."""
    # Combine texts and get a representative sample
    combined_text = "\n\n".join([f"Document: {text[:1500]}..." for text in document_texts])
    
    prompt = (
        f"Based on these financial document excerpts:\n\n{combined_text}\n\n"
        f"Determine the most appropriate audit framework for {audit_type or 'financial audit'}. "
        f"Examples include SA 700, AS 1, Ind AS 109, etc. "
        f"Provide the framework name and a brief explanation of why it's appropriate."
    )
    
    return call_groq_api(prompt)

def generate_suggested_questions(text):
    """Generate relevant audit questions based on document content."""
    # Limit text length for API
    sample_text = text[:4000]
    
    prompt = (
        f"Based on the following financial information:\n\n{sample_text}\n\n"
        "Generate 5 key questions an auditor should ask. "
        "Format each question as a numbered list (1., 2., etc.). "
        "Focus on potential risk areas, compliance concerns, and areas needing clarification."
    )
    
    questions = call_groq_api(prompt)
    return questions

def collect_company_info(state=None):
    """Collect and validate company KYC information."""
    if state is None:
        state = {}
    
    company_info = gr.Group(visible=True)
    with company_info:
        gr.Markdown("## Company Information (KYC)")
        company_name = gr.Textbox(label="Company Name", placeholder="Enter the company's legal name")
        company_id = gr.Textbox(label="Company Registration/CIN", placeholder="Enter company registration number")
        gst_id = gr.Textbox(label="GST Number", placeholder="Enter GST registration number")
        pan_number = gr.Textbox(label="PAN Number", placeholder="Enter company PAN")
        address = gr.Textbox(label="Registered Address", placeholder="Enter registered company address", lines=3)
        industry = gr.Dropdown(
            label="Industry Sector",
            choices=["Manufacturing", "Services", "Retail", "Technology", "Banking & Finance", "Healthcare", "Other"]
        )
        fiscal_year = gr.Textbox(label="Fiscal Year", placeholder="e.g., 2024-2025")
        submit_info = gr.Button("Save Company Information")
    
    return company_info, company_name, company_id, gst_id, pan_number, address, industry, fiscal_year, submit_info

def generate_audit_report_docx(audit_type, document_texts, framework):
    """Generate a professional DOCX audit report."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle with audit type
    subtitle = doc.add_heading(f'{audit_type} Assessment', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    from datetime import datetime
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    # Add executive summary
    doc.add_heading('Executive Summary', 1)
    summary_text = call_groq_api(
        f"Create an executive summary for an {audit_type} report based on these documents:\n\n" +
        "\n\n".join([text[:1500] + "..." for text in document_texts]) +
        "\n\nWrite a professional, concise executive summary (3-4 paragraphs)."
    )
    doc.add_paragraph(summary_text)
    
    # Add scope section
    doc.add_heading('Scope of Audit', 1)
    scope_text = call_groq_api(
        f"Create a scope section for an {audit_type} using framework {framework}. "
        "Describe what was covered in the audit, methodology used, and time period."
    )
    doc.add_paragraph(scope_text)
    
    # Add findings section
    doc.add_heading('Key Findings', 1)
    findings_text = call_groq_api(
        f"Generate key findings for an {audit_type} based on these documents:\n\n" +
        "\n\n".join([text[:1500] + "..." for text in document_texts]) +
        "\n\nCreate 3-5 significant findings with details."
        "\n\nProvide specific citations or references to the documents where applicable."
    )
    doc.add_paragraph(findings_text)
    
    # Add recommendations
    doc.add_heading('Recommendations', 1)
    recommendations_text = call_groq_api(
        f"Based on an {audit_type} audit with these findings:\n\n{findings_text}\n\n"
        "Provide 3-5 specific, actionable recommendations."
    )
    doc.add_paragraph(recommendations_text)
    
    # Add conclusion
    doc.add_heading('Conclusion', 1)
    conclusion_text = call_groq_api(
        f"Write a conclusion for an {audit_type} audit report that summarizes the overall assessment, "
        f"significance of findings, and next steps. Keep it professional and concise."
        f"Add final thoughts on the audit process and references to the documents reviewed."
    )
    doc.add_paragraph(conclusion_text)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file_path = temp_file.name
    temp_file.close()
    doc.save(temp_file_path)
    
    return temp_file_path

def initialize_chat():
    """Initialize the chat with a welcome message and audit options."""
    introduction = (
        "👋 Hello, I am Dabby Auditor - Your Auditing Copilot. I can assist you with various types of audits including:\n\n"
        "1. **Financial Statement Audits** - Examining financial records for accuracy and compliance\n"
        "2. **Compliance Audits** - Ensuring adherence to regulations like CARO, SA 230, IndAS, and GAAP\n"
        "3. **Internal Control Audits** - Evaluating the effectiveness of internal controls\n"
        "4. **Operational Audits** - Assessing efficiency and effectiveness of business operations\n"
        "5. **Tax Audits** - Reviewing tax compliance and identifying potential issues\n\n"
        "Feel free to tell me which type of audit you'd like to conduct, or upload your documents to get started."
    )
    
    history = []
    history.append({"role": "assistant", "content": introduction})
    return history

def generate_report_with_format(format_selection, history, state):
    """Generate an audit report in the selected format."""
    # Get document texts from state
    document_texts = state.get('document_texts', [])
    filenames = state.get('filenames', [])
    
    if not document_texts:
        return history, state, None, "No documents have been uploaded yet."
    
    try:
        # Map format selection to audit type
        format_to_audit_type = {
            "CARO Format": "Companies (Auditor's Report) Order",
            "SA 230 Format": "Standard on Auditing 230",
            "IndAS Format": "Indian Accounting Standards",
            "GAAP Format": "Generally Accepted Accounting Principles"
        }
        
        audit_type = format_to_audit_type.get(format_selection, "Financial Statement")
        
        # Determine appropriate framework if not already done
        if not state.get('framework'):
            state['framework'] = determine_audit_framework(document_texts, audit_type)
        framework = state['framework']
        
        # Generate and display methodology in chat
        methodology_prompt = (
            f"Create a detailed methodology for preparing a {audit_type} audit report following {framework} guidelines. "
            f"Explain the step-by-step process including:\n"
            f"1. Document analysis approach\n"
            f"2. Key areas examined\n"
            f"3. Analytical procedures used\n"
            f"4. How findings are evaluated\n"
            f"5. How recommendations are formulated\n"
            f"Format as a numbered research plan with clear headings and subpoints."
        )
        
        methodology = call_groq_api(methodology_prompt)
        
        # Add methodology to chat
        planning_message = (
            f"📋 **Audit Report Preparation Plan**\n\n"
            f"I'll now generate a {audit_type} audit report following {framework} guidelines. "
            f"Here's my methodology:\n\n{methodology}"
        )
        
        history.append({"role": "assistant", "content": planning_message})
        
        # Generate the report document
        report_file_path = generate_audit_report_docx(audit_type, document_texts, framework)
        
        # Success message
        message = "Report requested"
        response = (
            f"✅ **Report Generated Successfully**\n\n"
            f"I've prepared a comprehensive {audit_type} audit report based on "
            f"the uploaded documents ({', '.join(filenames)}). "
            f"The report follows {framework} guidelines and includes "
            f"executive summary, key findings, and recommendations. "
            f"You can download it using the button below."
        )
        
        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": response})
        
        return history, state, report_file_path, None
        
    except Exception as e:
        error_message = f"Error generating report: {str(e)}"
        print(error_message)
        return history, state, None, error_message

def audit_assistant(message, history, uploaded_files=None, state=None):
    # Initialize state if not present
    if state is None:
        state = {
            "document_texts": [],
            "filenames": [],
            "audit_type": None,
            "framework": None,
            "analysis_results": {},
            "files_processed": set(),  # Track already processed files
            "audit_stage": "introduction"  # Track the stage of the audit process
        }
    
    # Process only new uploaded files
    if uploaded_files:
        new_files = False
        for file in uploaded_files:
            file_path = file.name
            if file_path not in state.get("files_processed", set()):
                try:
                    text, filename = process_uploaded_file(file)
                    # Initialize lists if they don't exist
                    if "document_texts" not in state:
                        state["document_texts"] = []
                    if "filenames" not in state:
                        state["filenames"] = []
                    if "files_processed" not in state:
                        state["files_processed"] = set()
                                        
                    state["document_texts"].append(text)
                    state["filenames"].append(filename)
                    state["files_processed"].add(file_path)  # Mark as processed
                    new_files = True
                except Exception as e:
                    error_message = f"Error processing {file.name}: {str(e)}"
                    history.append({"role": "assistant", "content": error_message})
        
        # Only analyze if we have new files
        if new_files and state.get("document_texts", []):
            # Basic document analysis
            combined_text = "\n\n".join(state["document_texts"])
                        
            # Initialize analysis_results if it doesn't exist
            if "analysis_results" not in state:
                state["analysis_results"] = {}
                            
            # Only analyze once
            if "general" not in state.get("analysis_results", {}):
                analysis = analyze_document("Uploaded Documents", combined_text)
                state["analysis_results"]["general"] = analysis
                            
            # Determine audit type if not already set
            if not state.get("audit_type"):
                audit_type_prompt = (
                    f"Based on these financial documents, what type of audit would be most appropriate? "
                    f"Consider the content and purpose of the documents. Be specific but concise."
                )
                state["audit_type"] = call_groq_api(audit_type_prompt)
                            
            # Generate initial response for new files
            response = (
                f"I've analyzed the uploaded documents ({', '.join(state['filenames'])}). "
                f"Based on the content, this appears to be related to {state['audit_type']}.\n\n"
                f"Key observations:\n{state['analysis_results']['general'][:500]}...\n\n"
                f"You can ask specific questions about the documents or select a format to generate an audit report."
            )
                            
            history.append({"role": "assistant", "content": response})
    
    # Handle user messages
    if message:
            history.append({"role": "user", "content": message})
            
            # Check if user is selecting an audit type (if we're in the introduction stage)
            if not state.get("audit_type") or state.get("audit_stage") == "introduction":
                # Detect audit type from user message
                audit_type_prompt = (
                    f"The user message is: '{message}'\n\n"
                    f"Based on this message, determine which type of audit they want to conduct. "
                    f"Choose from: Financial Statement Audit, Compliance Audit, Internal Control Audit, "
                    f"Operational Audit, Tax Audit, or Other (specify). "
                    f"If they didn't specify an audit type, respond with 'Unspecified'."
                )
                detected_audit_type = call_groq_api(audit_type_prompt)
                
                if "unspecified" not in detected_audit_type.lower():
                    state["audit_type"] = detected_audit_type
                    state["audit_stage"] = "file_request"
                    
                    # Get required documents and best practices for this audit type
                    docs_prompt = (
                        f"For a {detected_audit_type}, list:\n"
                        f"1. The mandatory documents/files needed (5-7 items)\n"
                        f"2. Optional/supplementary documents that would be helpful (3-5 items)\n"
                        f"3. Best practices for conducting this type of audit (4-6 practices)\n\n"
                        f"Format as bullet points under clear headings."
                    )
                    audit_guidance = call_groq_api(docs_prompt)
                    
                    response = (
                        f"Brilliant choice! I'll help you conduct a {detected_audit_type}.\n\n"
                        f"{audit_guidance}\n\n"
                        f"Please upload the relevant documents, and we'll get started with the audit process."
                    )
                    history.append({"role": "assistant", "content": response})
                    return history, state
            
            # Check for report generation request
            if re.search(r'\b(generate|create|provide)\s+(a|an|the)?\s*report', message, re.IGNORECASE):
                response = (
                    "To generate a report, please select a format from the dropdown menu below. "
                    "Available formats include CARO, SA 230, IndAS, and GAAP."
                )
            elif not state.get("document_texts", []):  # Using .get() with default empty list
                response = (
                    "Please upload financial documents first so I can assist you with your audit. "
                    "I support PDF, Excel, Word, and CSV formats."
                )
            else:
                # Process general queries about the documents
                prompt = (
                    f"The user asks: {message}\n\n"
                    f"Based on these financial documents:\n" +
                    "\n\n".join([f"{filename}:\n{text[:1000]}..." for filename, text in zip(state.get("filenames", []), state.get("document_texts", []))]) +
                    "\n\nProvide a helpful, accurate response. If the information is not in the documents, say so clearly."
                )
                response = call_groq_api(prompt)
            
            history.append({"role": "assistant", "content": response})
        
            return history, state

# Set up the Gradio interface
with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("# AI Audit Assistant")
    gr.Markdown("Upload your financial documents and chat with the assistant to perform various types of audits.")
    
    with gr.Row():
        with gr.Column(scale=2):
            # File upload component 
            file_upload = gr.Files(
                label="Upload Financial Documents",
                type="filepath",
                file_count="multiple"
            )
            
            # Format selection dropdown
            format_dropdown = gr.Dropdown(
                label="Select Audit Report Format",
                choices=["CARO Format", "SA 230 Format", "IndAS Format", "GAAP Format"],
                value=None
            )
            
            # Report output
            report_output = gr.File(label="Download Generated Report")
            
            # Error messages
            error_output = gr.Textbox(label="", visible=False)
            
        with gr.Column(scale=3):
            # Chat interface
            chatbot = gr.Chatbot(
                label="Chatbot", 
                height=500,
                type="messages"  # Explicitly use the dictionary format
            )
            msg = gr.Textbox(label="Your message")
            clear = gr.Button("Clear Chat")
    
    # Create state
    state = gr.State({})
    
    # Initialize chat with introduction
    demo.load(
        initialize_chat,
        inputs=None,
        outputs=[chatbot],
        queue=False
    )
    
    # Handle chat interactions
    msg.submit(
        audit_assistant,
        inputs=[msg, chatbot, file_upload, state],
        outputs=[chatbot, state],
        queue=False
    )
    
    # Handle file uploads
    file_upload.change(
        audit_assistant,
        inputs=[gr.Textbox(value="", visible=False), chatbot, file_upload, state],
        outputs=[chatbot, state],
        queue=False
    )
    
    # Handle report generation
    format_dropdown.change(
        generate_report_with_format,
        inputs=[format_dropdown, chatbot, state],
        outputs=[chatbot, state, report_output, error_output],
        queue=False
    )
    
    # Clear chat functionality - also reset state
    def clear_chat():
        initial_history = initialize_chat()
        return initial_history, {}
    
    clear.click(clear_chat, None, [chatbot, state], queue=False)

# Launch the Gradio app
if __name__ == "__main__":
    demo.launch(share=True)
